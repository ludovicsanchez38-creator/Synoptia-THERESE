"""
THERESE v2 - Conversion Markdown -> DOCX déterministe (BUG-135).

Chemin d'export des documents de l'Atelier documentaire et des conversations.
Contrairement au pipeline LLM (CodeGenSkill/DocxSkill), ce convertisseur :
- ne supprime JAMAIS de contenu : les blocs ``` clôturés sont rendus en
  préformaté, une fence orpheline ne tronque pas la suite du document ;
- n'exécute jamais de code présent dans le document (un bloc ```python```
  est du contenu, pas un générateur) ;
- pose la langue de correction fr-FR (docDefaults + style Normal), au lieu
  du en-US hérité du template python-docx.
"""

import re
from pathlib import Path

from app.services.export_profile import ExportProfile
from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

# Segments inline **gras**, *italique*, `code` - tout le reste (y compris les
# * ou ` non appariés) est ajouté tel quel, jamais jeté.
_INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|\*[^*\n]+?\*|`[^`\n]+?`)")
_TABLE_SEPARATOR = re.compile(r"^\|?[\s\-:|]+\|?\s*$")
_HR_PATTERN = re.compile(r"^(-{3,}|\*{3,}|_{3,})\s*$")


def render_markdown_docx(
    markdown: str, output_path: Path, profile: ExportProfile | None = None
) -> None:
    """Rend un document Markdown complet en .docx. Le profil d'export
    (chantier 5) pilote langue, polices, couleurs, footer et marges - ses
    DÉFAUTS reproduisent la charte Synoptia historique (fr-FR, Calibri 11)."""
    profile = profile or ExportProfile()
    doc = Document()
    _setup_styles(doc, profile)
    _set_language(doc, profile.language)
    _apply_margins(doc, profile)
    _render_blocks(doc, markdown, profile)
    _add_footer(doc, profile)
    doc.save(str(output_path))


def _color(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value.lstrip("#"))


def _apply_margins(doc: DocumentObject, profile: ExportProfile) -> None:
    m = profile.margins_cm
    for section in doc.sections:
        section.top_margin = Cm(m.top)
        section.bottom_margin = Cm(m.bottom)
        section.left_margin = Cm(m.left)
        section.right_margin = Cm(m.right)


def _set_language(doc: DocumentObject, lang: str) -> None:
    """Pose w:lang sur les docDefaults ET le style Normal (correction Word)."""
    rpr_defaults = doc.styles.element.findall(
        f"{qn('w:docDefaults')}/{qn('w:rPrDefault')}/{qn('w:rPr')}"
    )
    targets = list(rpr_defaults)
    targets.append(doc.styles["Normal"].element.get_or_add_rPr())
    for rpr in targets:
        for existing in rpr.findall(qn("w:lang")):
            rpr.remove(existing)
        lang_el = rpr.makeelement(qn("w:lang"), {})
        lang_el.set(qn("w:val"), lang)
        lang_el.set(qn("w:eastAsia"), lang)
        rpr.append(lang_el)


def _render_blocks(doc: DocumentObject, markdown: str, profile: ExportProfile) -> None:
    """Machine à états ligne par ligne - aucune branche ne jette de contenu."""
    in_code = False
    first_h1_done = False
    table_buffer: list[str] = []

    for raw in markdown.split("\n"):
        line = raw.strip()

        if in_code:
            if line.startswith("```"):
                in_code = False
            else:
                _add_preformatted(doc, raw)
            continue

        if not line.startswith("|"):
            _flush_table(doc, table_buffer)

        if line.startswith("```"):
            # Fence ouvrante : le contenu sera rendu préformaté. Si elle n'est
            # jamais refermée, TOUT le reste sort en préformaté (zéro perte).
            in_code = True
            continue

        if not line:
            continue

        if line.startswith("|"):
            table_buffer.append(line)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            if first_h1_done:
                doc.add_heading(line[2:], level=1)
            else:
                # Premier # = titre principal du document (style hero), les
                # suivants redeviennent des Heading 1 ordinaires.
                para = doc.add_heading(line[2:], level=0)
                _style_title(para, profile)
                first_h1_done = True
        elif _HR_PATTERN.match(line):
            continue  # séparateur horizontal : pas de texte littéral
        elif line.startswith("> "):
            para = doc.add_paragraph()
            _add_formatted_text(para, line[2:])
            for run in para.runs:
                run.italic = True
        elif line.startswith("- ") or line.startswith("* "):
            para = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(para, line[2:])
        elif re.match(r"^\d+\.\s", line):
            para = doc.add_paragraph(style="List Number")
            _add_formatted_text(para, re.sub(r"^\d+\.\s", "", line))
        else:
            para = doc.add_paragraph()
            _add_formatted_text(para, line)

    _flush_table(doc, table_buffer)


def _flush_table(doc: DocumentObject, buffer: list[str]) -> None:
    """Rend un groupe de lignes `|...|`. Vrai tableau Word si la 2e ligne est
    un séparateur d'en-tête Markdown, sinon repli en paragraphes (les cellules
    jointes par tabulation - comportement historique, sans perte)."""
    if not buffer:
        return
    rows = list(buffer)
    buffer.clear()

    if len(rows) >= 2 and _TABLE_SEPARATOR.match(rows[1]):
        header = _split_row(rows[0])
        body = [_split_row(r) for r in rows[2:] if not _TABLE_SEPARATOR.match(r)]
        # Largeur = ligne la plus large (revue 10/07 : borner au nombre de
        # colonnes de l'en-tête perdait silencieusement les cellules
        # excédentaires du corps - zéro perte, quitte à padder l'en-tête).
        ncols = max(len(header), *(len(r) for r in body)) if body else len(header)
        table = doc.add_table(rows=1, cols=ncols)
        table.style = "Table Grid"
        for j, cell_text in enumerate(header):
            run = table.rows[0].cells[j].paragraphs[0].add_run(cell_text)
            run.bold = True
        for row_cells in body:
            row = table.add_row()
            for j in range(ncols):
                text = row_cells[j] if j < len(row_cells) else ""
                _add_formatted_text(row.cells[j].paragraphs[0], text)
        return

    for r in rows:
        if _TABLE_SEPARATOR.match(r):
            continue
        para = doc.add_paragraph()
        para.add_run("\t".join(_split_row(r)))


def _split_row(line: str) -> list[str]:
    """Cellules d'une ligne de tableau - tolère l'absence de pipe fermant
    (l'ancien split('|')[1:-1] perdait la dernière cellule) et respecte les
    pipes échappés `\\|` dans une cellule (revue 10/07)."""
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|") and not line.endswith("\\|"):
        line = line[:-1]
    cells = re.split(r"(?<!\\)\|", line)
    return [c.strip().replace("\\|", "|") for c in cells]


def _add_preformatted(doc: DocumentObject, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def _add_formatted_text(para: Paragraph, text: str) -> None:
    """Inline **gras**/*italique*/`code`. Les segments hors motifs (dont les
    * et ` non appariés) sont ajoutés tels quels : rien n'est jeté."""
    pos = 0
    for m in _INLINE_PATTERN.finditer(text):
        if m.start() > pos:
            para.add_run(text[pos : m.start()])
        part = m.group(0)
        if part.startswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*"):
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            run = para.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        pos = m.end()
    if pos < len(text):
        para.add_run(text[pos:])


def _setup_styles(doc: DocumentObject, profile: ExportProfile) -> None:
    """Styles pilotés par le profil d'export (défauts = charte Synoptia)."""
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = profile.body_font
    normal.font.size = Pt(profile.body_size_pt)
    normal.font.color.rgb = _color(profile.body_color)
    normal.paragraph_format.space_after = Pt(10)
    normal.paragraph_format.line_spacing = 1.15

    heading_specs = {
        "Heading 1": (24, profile.heading_color, 24, 12),
        "Heading 2": (18, profile.h2_color, 18, 8),
        "Heading 3": (14, profile.heading_color, 14, 6),
        "Heading 4": (12, profile.heading_color, 12, 6),
    }
    for name, (size, color, before, after) in heading_specs.items():
        if name in styles:
            style = styles[name]
            style.font.name = profile.heading_font
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = _color(color)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)


def _style_title(para: Paragraph, profile: ExportProfile) -> None:
    for run in para.runs:
        run.font.name = profile.heading_font
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = _color(profile.title_color)


def _add_footer(doc: DocumentObject, profile: ExportProfile) -> None:
    if not profile.footer_text:
        return
    section = doc.sections[-1]
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(profile.footer_text)
    run.font.size = Pt(9)
