"""
THÉRÈSE v2 - Word Document Generator Skill

Génère des documents Word (.docx) avec le style Synoptïa.
"""

import logging
import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.services.skills.base import BaseSkill, FileFormat, SkillParams, SkillResult

logger = logging.getLogger(__name__)


# Palette Synoptïa
SYNOPTIA_COLORS = {
    "background": RGBColor(0x0B, 0x12, 0x26),
    "text": RGBColor(0xE6, 0xED, 0xF7),
    "muted": RGBColor(0xA9, 0xB8, 0xD8),
    "primary": RGBColor(0x17, 0x33, 0xA6),
    "accent_cyan": RGBColor(0x22, 0xD3, 0xEE),
    "accent_magenta": RGBColor(0xE1, 0x1D, 0x8D),
    # Couleurs claires pour documents imprimables
    "heading": RGBColor(0x0F, 0x1E, 0x6D),
    "body": RGBColor(0x1A, 0x1A, 0x2E),
}


class DocxSkill(BaseSkill):
    """
    Skill de génération de documents Word.

    Crée des documents .docx professionnels avec le style Synoptïa.
    """

    skill_id = "docx-pro"
    name = "Document Word Professionnel"
    description = "Génère un document Word structuré avec le style Synoptïa"
    output_format = FileFormat.DOCX

    def __init__(self, output_dir: Path):
        super().__init__(output_dir)

    async def execute(self, params: SkillParams) -> SkillResult:
        """
        Génère un document Word à partir du contenu.

        Args:
            params: Paramètres incluant titre et contenu

        Returns:
            Résultat avec chemin vers le fichier généré
        """
        file_id = self.generate_file_id()
        output_path = self.get_output_path(file_id, params.title)

        # Créer le document
        doc = Document()

        # Appliquer les styles Synoptïa
        self._setup_styles(doc)

        # Ajouter le titre principal
        title_para = doc.add_heading(params.title, level=0)
        self._style_title(title_para)

        # Parser et ajouter le contenu
        self._add_content(doc, params.content)

        # Ajouter le footer Synoptïa
        self._add_footer(doc)

        # Sauvegarder
        doc.save(str(output_path))

        # Calculer la taille
        file_size = output_path.stat().st_size

        logger.info(f"Generated DOCX: {output_path} ({file_size} bytes)")

        return SkillResult(
            file_id=file_id,
            file_path=output_path,
            file_name=output_path.name,
            file_size=file_size,
            mime_type=self.get_mime_type(),
            format=self.output_format,
        )

    def get_system_prompt_addition(self) -> str:
        """Instructions pour le LLM pour générer du contenu Word."""
        return """
## Instructions pour génération de document Word

Génère le contenu du document en format Markdown structuré :

1. **Titres** : Utilise les niveaux de titre Markdown (## pour h2, ### pour h3, etc.)
2. **Paragraphes** : Sépare les paragraphes par des lignes vides
3. **Listes** : Utilise - pour les listes à puces, 1. pour les listes numérotées
4. **Mise en forme** : **gras** pour les mots importants, *italique* pour l'emphase
5. **Tableaux** : Format Markdown standard si nécessaire

Structure recommandée :
- Introduction (contexte, objectif)
- Corps (sections logiques)
- Conclusion (résumé, prochaines étapes)

Style : professionnel, concis, orienté action.
"""

    def _setup_styles(self, doc: Document) -> None:
        """Configure les styles du document."""
        styles = doc.styles

        # Style Normal
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        normal.font.color.rgb = SYNOPTIA_COLORS["body"]
        normal.paragraph_format.space_after = Pt(10)
        normal.paragraph_format.line_spacing = 1.15

        # Style Heading 1
        if "Heading 1" in styles:
            h1 = styles["Heading 1"]
            h1.font.name = "Outfit"
            h1.font.size = Pt(24)
            h1.font.bold = True
            h1.font.color.rgb = SYNOPTIA_COLORS["heading"]
            h1.paragraph_format.space_before = Pt(24)
            h1.paragraph_format.space_after = Pt(12)

        # Style Heading 2
        if "Heading 2" in styles:
            h2 = styles["Heading 2"]
            h2.font.name = "Outfit"
            h2.font.size = Pt(18)
            h2.font.bold = True
            h2.font.color.rgb = SYNOPTIA_COLORS["primary"]
            h2.paragraph_format.space_before = Pt(18)
            h2.paragraph_format.space_after = Pt(8)

        # Style Heading 3
        if "Heading 3" in styles:
            h3 = styles["Heading 3"]
            h3.font.name = "Outfit"
            h3.font.size = Pt(14)
            h3.font.bold = True
            h3.font.color.rgb = SYNOPTIA_COLORS["heading"]
            h3.paragraph_format.space_before = Pt(14)
            h3.paragraph_format.space_after = Pt(6)

    def _style_title(self, para) -> None:
        """Applique le style au titre principal."""
        for run in para.runs:
            run.font.name = "Outfit"
            run.font.size = Pt(28)
            run.font.bold = True
            run.font.color.rgb = SYNOPTIA_COLORS["heading"]

    def _add_content(self, doc: Document, content: str) -> None:
        """
        Parse le contenu Markdown et l'ajoute au document.

        Args:
            doc: Document Word
            content: Contenu en format Markdown
        """
        lines = content.split('\n')
        current_list_type = None  # 'bullet' ou 'numbered'
        list_number = 0

        for line in lines:
            line = line.strip()
            if not line:
                current_list_type = None
                list_number = 0
                continue

            # Headings
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
                current_list_type = None
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
                current_list_type = None
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
                current_list_type = None
            # Listes à puces
            elif line.startswith('- ') or line.startswith('* '):
                text = line[2:]
                para = doc.add_paragraph(style='List Bullet')
                self._add_formatted_text(para, text)
                current_list_type = 'bullet'
            # Listes numérotées
            elif re.match(r'^\d+\.\s', line):
                text = re.sub(r'^\d+\.\s', '', line)
                para = doc.add_paragraph(style='List Number')
                self._add_formatted_text(para, text)
                current_list_type = 'numbered'
                list_number += 1
            # Tableaux (simplifié)
            elif line.startswith('|'):
                # Ignorer les lignes de séparation
                if not re.match(r'^\|[\s\-:|]+\|$', line):
                    cells = [c.strip() for c in line.split('|')[1:-1]]
                    if cells:
                        para = doc.add_paragraph()
                        para.add_run('\t'.join(cells))
            # Paragraphe normal
            else:
                para = doc.add_paragraph()
                self._add_formatted_text(para, line)
                current_list_type = None

    def _add_formatted_text(self, para, text: str) -> None:
        """
        Ajoute du texte formaté (gras, italique) à un paragraphe.

        Args:
            para: Paragraphe Word
            text: Texte avec formatage Markdown
        """
        # Pattern pour **gras**, *italique*, `code`
        pattern = r'(\*\*.*?\*\*|\*.*?\*|`.*?`|[^*`]+)'
        parts = re.findall(pattern, text)

        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = para.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = para.add_run(part[1:-1])
                run.font.name = "Consolas"
                run.font.size = Pt(10)
            else:
                para.add_run(part)

    def _add_footer(self, doc: Document) -> None:
        """Ajoute un footer Synoptïa au document."""
        section = doc.sections[-1]
        footer = section.footer
        footer.is_linked_to_previous = False

        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = para.add_run("Généré par THÉRÈSE - Synoptïa")
        run.font.size = Pt(9)
        run.font.color.rgb = SYNOPTIA_COLORS["muted"]
        run.font.italic = True
