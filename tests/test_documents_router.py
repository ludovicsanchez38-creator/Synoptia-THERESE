"""
THÉRÈSE v2 - Documents Router Tests

Tests du CRUD pur de l'atelier documentaire (documents, sections, pistes).
Le point le plus critique est l'invariant de complétude à la réorganisation
de la trame : un écart entre les ids reçus et les ids en base doit renvoyer
409 SANS AUCUNE écriture (garde-fou anti-perte de données).
"""

import json
from datetime import datetime
from unittest.mock import AsyncMock, patch

import pytest
from httpx import AsyncClient


async def _create_document(client: AsyncClient, title: str = "Proposition ISOCUBE", brief: str = "Brief du besoin") -> dict:
    response = await client.post("/api/documents", json={"title": title, "brief": brief})
    assert response.status_code == 200, response.text
    return response.json()


async def _create_section(
    client: AsyncClient, document_id: str, title: str, order: float, depth: int = 0
) -> dict:
    response = await client.post(
        f"/api/documents/{document_id}/sections",
        json={"title": title, "brief": "", "order": order, "depth": depth},
    )
    assert response.status_code == 200, response.text
    return response.json()


async def _document_updated_at(client: AsyncClient, document_id: str) -> datetime:
    response = await client.get(f"/api/documents/{document_id}")
    assert response.status_code == 200, response.text
    return datetime.fromisoformat(response.json()["updated_at"])


class TestDocumentsCRUD:
    """CRUD nominal des documents."""

    @pytest.mark.asyncio
    async def test_create_document(self, client: AsyncClient):
        """POST /api/documents crée un document vide (0 section)."""
        response = await client.post(
            "/api/documents",
            json={"title": "Proposition ISOCUBE", "brief": "Formation Claude 2h30"},
        )

        assert response.status_code == 200
        doc = response.json()
        assert doc["title"] == "Proposition ISOCUBE"
        assert doc["brief"] == "Formation Claude 2h30"
        assert doc["status"] == "en_cours"
        assert doc["sections_total"] == 0
        assert doc["sections_validees"] == 0
        assert "id" in doc
        assert "created_at" in doc
        assert "updated_at" in doc

    @pytest.mark.asyncio
    async def test_list_documents_empty(self, client: AsyncClient):
        """GET /api/documents retourne une liste vide au départ."""
        response = await client.get("/api/documents")

        assert response.status_code == 200
        assert response.json() == []

    @pytest.mark.asyncio
    async def test_list_documents_with_sections_count(self, client: AsyncClient):
        """GET /api/documents agrège correctement le compteur de sections."""
        doc = await _create_document(client)
        await _create_section(client, doc["id"], "Section 1", order=0.0)
        await _create_section(client, doc["id"], "Section 2", order=1.0)

        response = await client.get("/api/documents")

        assert response.status_code == 200
        docs = response.json()
        assert len(docs) == 1
        assert docs[0]["id"] == doc["id"]
        assert docs[0]["sections_total"] == 2
        assert docs[0]["sections_validees"] == 0

    @pytest.mark.asyncio
    async def test_get_document_with_sections_and_pistes(self, client: AsyncClient):
        """GET /api/documents/{id} renvoie sections triées par ordre + pistes."""
        doc = await _create_document(client)
        await _create_section(client, doc["id"], "Section B", order=2.0)
        await _create_section(client, doc["id"], "Section A", order=1.0)
        await client.post(
            f"/api/documents/{doc['id']}/pistes",
            json={"texte": "Creuser le volet OPCO"},
        )

        response = await client.get(f"/api/documents/{doc['id']}")

        assert response.status_code == 200
        detail = response.json()
        assert detail["id"] == doc["id"]
        assert [s["title"] for s in detail["sections"]] == ["Section A", "Section B"]
        assert detail["sections_total"] == 2
        assert len(detail["pistes"]) == 1
        assert detail["pistes"][0]["texte"] == "Creuser le volet OPCO"
        assert detail["pistes"][0]["status"] == "nouvelle"

    @pytest.mark.asyncio
    async def test_get_document_not_found(self, client: AsyncClient):
        """GET /api/documents/{id} renvoie 404 si le document n'existe pas."""
        response = await client.get("/api/documents/id-inexistant")
        assert response.status_code == 404

    @pytest.mark.asyncio
    async def test_delete_document_cascades_sections_and_pistes(self, client: AsyncClient):
        """DELETE /api/documents/{id} supprime le document, ses sections et ses pistes."""
        doc = await _create_document(client)
        section = await _create_section(client, doc["id"], "Section 1", order=0.0)
        piste_resp = await client.post(
            f"/api/documents/{doc['id']}/pistes",
            json={"texte": "Piste à explorer"},
        )
        piste = piste_resp.json()

        response = await client.delete(f"/api/documents/{doc['id']}")

        assert response.status_code == 200
        assert response.json()["success"] is True

        # Le document a disparu
        assert (await client.get(f"/api/documents/{doc['id']}")).status_code == 404

        # La section a été cascadée (PATCH sur son id renvoie 404)
        section_check = await client.patch(
            f"/api/documents/sections/{section['id']}",
            json={"title": "Fantôme"},
        )
        assert section_check.status_code == 404

        # La piste a été cascadée (PATCH sur son id renvoie 404)
        piste_check = await client.patch(
            f"/api/documents/pistes/{piste['id']}",
            json={"status": "exploree"},
        )
        assert piste_check.status_code == 404

    @pytest.mark.asyncio
    async def test_delete_document_not_found(self, client: AsyncClient):
        """DELETE /api/documents/{id} renvoie 404 si le document n'existe pas."""
        response = await client.delete("/api/documents/id-inexistant")
        assert response.status_code == 404


class TestSections:
    """Création et mise à jour des sections."""

    @pytest.mark.asyncio
    async def test_create_section(self, client: AsyncClient):
        """POST /api/documents/{id}/sections crée une section manuelle."""
        doc = await _create_document(client)

        response = await client.post(
            f"/api/documents/{doc['id']}/sections",
            json={"title": "Contexte", "brief": "Poser le décor", "order": 0.0, "depth": 0},
        )

        assert response.status_code == 200
        section = response.json()
        assert section["title"] == "Contexte"
        assert section["brief"] == "Poser le décor"
        assert section["status"] == "vide"
        assert section["orphan"] is False
        assert section["document_id"] == doc["id"]

    @pytest.mark.asyncio
    async def test_create_section_document_not_found(self, client: AsyncClient):
        """POST .../sections renvoie 404 si le document n'existe pas."""
        response = await client.post(
            "/api/documents/id-inexistant/sections",
            json={"title": "Contexte", "order": 0.0},
        )
        assert response.status_code == 404

    @pytest.mark.asyncio
    async def test_create_section_bumps_document_updated_at(self, client: AsyncClient):
        """Créer une section fait remonter le document dans la liste triée
        par updated_at desc."""
        doc = await _create_document(client)
        before = await _document_updated_at(client, doc["id"])

        await _create_section(client, doc["id"], "Nouvelle section", order=0.0)

        after = await _document_updated_at(client, doc["id"])
        assert after > before

    @pytest.mark.asyncio
    async def test_patch_content_sets_status_brouillon(self, client: AsyncClient):
        """PATCH content sur une section 'vide' fait passer le statut à 'brouillon'."""
        doc = await _create_document(client)
        section = await _create_section(client, doc["id"], "Section 1", order=0.0)
        assert section["status"] == "vide"

        response = await client.patch(
            f"/api/documents/sections/{section['id']}",
            json={"content": "Premier jet de rédaction."},
        )

        assert response.status_code == 200
        updated = response.json()
        assert updated["content"] == "Premier jet de rédaction."
        assert updated["status"] == "brouillon"

    @pytest.mark.asyncio
    async def test_patch_title_only_does_not_change_status(self, client: AsyncClient):
        """PATCH sans toucher au contenu ne fait pas sortir une section de 'vide'."""
        doc = await _create_document(client)
        section = await _create_section(client, doc["id"], "Section 1", order=0.0)

        response = await client.patch(
            f"/api/documents/sections/{section['id']}",
            json={"title": "Nouveau titre"},
        )

        assert response.status_code == 200
        updated = response.json()
        assert updated["title"] == "Nouveau titre"
        assert updated["status"] == "vide"

    @pytest.mark.asyncio
    async def test_patch_section_not_found(self, client: AsyncClient):
        """PATCH sur une section inexistante renvoie 404."""
        response = await client.patch(
            "/api/documents/sections/id-inexistant",
            json={"title": "X"},
        )
        assert response.status_code == 404

    @pytest.mark.asyncio
    async def test_patch_title_null_does_not_500_and_keeps_title(self, client: AsyncClient):
        """PATCH {"title": null} : un null explicite ne doit jamais tenter
        d'écrire NULL sur une colonne non-nullable (500 IntegrityError) - il
        est traité comme une absence de changement, comme un champ omis."""
        doc = await _create_document(client)
        section = await _create_section(client, doc["id"], "Titre original", order=0.0)

        response = await client.patch(
            f"/api/documents/sections/{section['id']}",
            json={"title": None},
        )

        assert response.status_code == 200, response.text
        updated = response.json()
        assert updated["title"] == "Titre original"

    @pytest.mark.asyncio
    async def test_update_section_bumps_document_updated_at(self, client: AsyncClient):
        """Une section retouchée fait remonter son document (liste triée par
        updated_at desc) - PATCH doit bumper `Document.updated_at`, pas
        seulement `DocumentSection.updated_at`."""
        doc = await _create_document(client)
        section = await _create_section(client, doc["id"], "Section 1", order=0.0)
        before = await _document_updated_at(client, doc["id"])

        response = await client.patch(
            f"/api/documents/sections/{section['id']}",
            json={"content": "Un contenu qui fait vivre le document."},
        )
        assert response.status_code == 200, response.text

        after = await _document_updated_at(client, doc["id"])
        assert after > before


class TestOutlineGeneration:
    """POST /api/documents/{id}/outline - génération de trame via le LLM (mocké)."""

    @pytest.mark.asyncio
    async def test_outline_nominal_creates_sorted_sections(self, client: AsyncClient):
        """Une réponse LLM lisible crée les sections triées, order 10/20/30, depths respectés."""
        doc = await _create_document(client, title="Proposition ISOCUBE", brief="Formation Claude 2h30")

        raw_outline = json.dumps(
            [
                {"title": "Contexte", "brief": "Poser le décor", "depth": 0},
                {"title": "Détail", "brief": "Approfondir un point", "depth": 1},
                {"title": "Conclusion", "brief": "Résumer et conclure", "depth": 0},
            ]
        )

        with patch(
            "app.services.llm.LLMService.generate_content",
            new_callable=AsyncMock,
            return_value=raw_outline,
        ) as mock_generate:
            response = await client.post(f"/api/documents/{doc['id']}/outline")

        assert response.status_code == 200, response.text
        mock_generate.assert_called_once()

        sections = response.json()
        assert len(sections) == 3
        assert [s["title"] for s in sections] == ["Contexte", "Détail", "Conclusion"]
        assert [s["order"] for s in sections] == [10.0, 20.0, 30.0]
        assert [s["depth"] for s in sections] == [0, 1, 0]
        assert all(s["status"] == "vide" for s in sections)
        assert all(s["document_id"] == doc["id"] for s in sections)
        # Les ids sont stables : uniques et posés une fois pour toutes.
        assert len({s["id"] for s in sections}) == 3

        # Persisté en base : GET retrouve la trame triée.
        detail = await client.get(f"/api/documents/{doc['id']}")
        assert [s["title"] for s in detail.json()["sections"]] == [
            "Contexte",
            "Détail",
            "Conclusion",
        ]

    @pytest.mark.asyncio
    async def test_outline_unreadable_llm_response_returns_502_and_creates_nothing(
        self, client: AsyncClient
    ):
        """Une réponse LLM illisible -> 502 'trame illisible' et AUCUNE section créée."""
        doc = await _create_document(client)

        with patch(
            "app.services.llm.LLMService.generate_content",
            new_callable=AsyncMock,
            return_value="ceci n'est pas du JSON du tout",
        ):
            response = await client.post(f"/api/documents/{doc['id']}/outline")

        assert response.status_code == 502
        assert "trame illisible" in response.json()["message"].lower()

        detail = await client.get(f"/api/documents/{doc['id']}")
        assert detail.json()["sections"] == []
        assert detail.json()["sections_total"] == 0

    @pytest.mark.asyncio
    async def test_outline_existing_non_empty_sections_returns_409_without_llm_call(
        self, client: AsyncClient
    ):
        """Un document avec une section déjà rédigée -> 409, aucun appel LLM."""
        doc = await _create_document(client)
        section = await _create_section(client, doc["id"], "Section 1", order=0.0)
        await client.patch(
            f"/api/documents/sections/{section['id']}",
            json={"content": "Contenu déjà rédigé."},
        )

        with patch(
            "app.services.llm.LLMService.generate_content",
            new_callable=AsyncMock,
        ) as mock_generate:
            response = await client.post(f"/api/documents/{doc['id']}/outline")

        assert response.status_code == 409
        mock_generate.assert_not_called()

        # AUCUNE écriture : la section rédigée d'origine est intacte, pas de nouvelle section.
        detail = await client.get(f"/api/documents/{doc['id']}")
        assert detail.json()["sections_total"] == 1
        assert detail.json()["sections"][0]["content"] == "Contenu déjà rédigé."

    @pytest.mark.asyncio
    async def test_outline_document_not_found(self, client: AsyncClient):
        """POST .../outline renvoie 404 si le document n'existe pas."""
        with patch(
            "app.services.llm.LLMService.generate_content",
            new_callable=AsyncMock,
        ) as mock_generate:
            response = await client.post("/api/documents/id-inexistant/outline")

        assert response.status_code == 404
        mock_generate.assert_not_called()

    @pytest.mark.asyncio
    async def test_outline_double_appel_remplace_sans_doublons(self, client: AsyncClient):
        """Durcissement issu de la revue B2 : double-clic (ou re-génération
        après une 1re trame jamais retouchée) - la trame reste entièrement
        vide entre les deux appels, donc le garde-fou 409 ne se déclenche
        pas. Le 2e appel doit REMPLACER la 1re trame, jamais s'additionner
        (pas de doublons d'order 10/20)."""
        doc = await _create_document(client)

        raw_outline = json.dumps(
            [
                {"title": "Contexte", "brief": "Poser le décor", "depth": 0},
                {"title": "Conclusion", "brief": "Résumer et conclure", "depth": 0},
            ]
        )

        with patch(
            "app.services.llm.LLMService.generate_content",
            new_callable=AsyncMock,
            return_value=raw_outline,
        ):
            first = await client.post(f"/api/documents/{doc['id']}/outline")
            assert first.status_code == 200, first.text
            second = await client.post(f"/api/documents/{doc['id']}/outline")
            assert second.status_code == 200, second.text

        detail = await client.get(f"/api/documents/{doc['id']}")
        sections = detail.json()["sections"]
        assert len(sections) == 2
        assert [s["order"] for s in sections] == [10.0, 20.0]
        assert [s["title"] for s in sections] == ["Contexte", "Conclusion"]


class TestSectionsReorderInvariant:
    """L'invariant de complétude - le cœur du design de l'atelier documentaire."""

    @pytest.mark.asyncio
    async def test_reorder_missing_id_returns_409_and_writes_nothing(self, client: AsyncClient):
        """Un id manquant dans la réorganisation -> 409 ET aucune écriture."""
        doc = await _create_document(client)
        section_a = await _create_section(client, doc["id"], "Section A", order=0.0, depth=0)
        section_b = await _create_section(client, doc["id"], "Section B", order=1.0, depth=0)

        # Réorganisation incomplète : section_b est absente de la requête.
        response = await client.post(
            f"/api/documents/{doc['id']}/sections/reorder",
            json={"items": [{"id": section_a["id"], "order": 5.0, "depth": 1}]},
        )

        assert response.status_code == 409
        body = response.json()
        assert section_b["id"] in body["missing_ids"]
        assert body["unknown_ids"] == []

        # AUCUNE écriture : l'ordre et la profondeur d'origine sont intacts.
        check = await client.get(f"/api/documents/{doc['id']}")
        sections_by_id = {s["id"]: s for s in check.json()["sections"]}
        assert sections_by_id[section_a["id"]]["order"] == 0.0
        assert sections_by_id[section_a["id"]]["depth"] == 0
        assert sections_by_id[section_b["id"]]["order"] == 1.0
        assert sections_by_id[section_b["id"]]["depth"] == 0

    @pytest.mark.asyncio
    async def test_reorder_unknown_id_returns_409_and_writes_nothing(self, client: AsyncClient):
        """Un id inconnu (n'existant pas en base) dans la réorganisation -> 409."""
        doc = await _create_document(client)
        section_a = await _create_section(client, doc["id"], "Section A", order=0.0)

        response = await client.post(
            f"/api/documents/{doc['id']}/sections/reorder",
            json={
                "items": [
                    {"id": section_a["id"], "order": 0.0, "depth": 0},
                    {"id": "id-fantome", "order": 1.0, "depth": 0},
                ]
            },
        )

        assert response.status_code == 409
        body = response.json()
        assert body["missing_ids"] == []
        assert body["unknown_ids"] == ["id-fantome"]

        # AUCUNE écriture : la section existante n'a pas bougé.
        check = await client.get(f"/api/documents/{doc['id']}")
        assert check.json()["sections"][0]["order"] == 0.0

    @pytest.mark.asyncio
    async def test_reorder_valid_reorders_sections(self, client: AsyncClient):
        """Une réorganisation complète (ensemble d'ids exact) réordonne réellement."""
        doc = await _create_document(client)
        section_a = await _create_section(client, doc["id"], "Section A", order=0.0, depth=0)
        section_b = await _create_section(client, doc["id"], "Section B", order=1.0, depth=0)

        response = await client.post(
            f"/api/documents/{doc['id']}/sections/reorder",
            json={
                "items": [
                    {"id": section_a["id"], "order": 10.0, "depth": 1},
                    {"id": section_b["id"], "order": 5.0, "depth": 0},
                ]
            },
        )

        assert response.status_code == 200
        reordered = response.json()
        # Réponse triée par ordre croissant : B (5.0) puis A (10.0).
        assert [s["id"] for s in reordered] == [section_b["id"], section_a["id"]]

        # Persisté en base.
        check = await client.get(f"/api/documents/{doc['id']}")
        sections_by_id = {s["id"]: s for s in check.json()["sections"]}
        assert sections_by_id[section_a["id"]]["order"] == 10.0
        assert sections_by_id[section_a["id"]]["depth"] == 1
        assert sections_by_id[section_b["id"]]["order"] == 5.0

    @pytest.mark.asyncio
    async def test_reorder_document_not_found(self, client: AsyncClient):
        """POST .../sections/reorder renvoie 404 si le document n'existe pas."""
        response = await client.post(
            "/api/documents/id-inexistant/sections/reorder",
            json={"items": []},
        )
        assert response.status_code == 404

    @pytest.mark.asyncio
    async def test_reorder_valid_bumps_document_updated_at(self, client: AsyncClient):
        """Une réorganisation réussie fait remonter le document dans la
        liste triée par updated_at desc."""
        doc = await _create_document(client)
        section_a = await _create_section(client, doc["id"], "Section A", order=0.0)
        before = await _document_updated_at(client, doc["id"])

        response = await client.post(
            f"/api/documents/{doc['id']}/sections/reorder",
            json={"items": [{"id": section_a["id"], "order": 5.0, "depth": 0}]},
        )
        assert response.status_code == 200, response.text

        after = await _document_updated_at(client, doc["id"])
        assert after > before


class TestPistes:
    """Liste, création manuelle et changement de statut des pistes."""

    @pytest.mark.asyncio
    async def test_create_and_list_pistes(self, client: AsyncClient):
        """POST puis GET .../pistes crée et liste les pistes d'un document."""
        doc = await _create_document(client)

        create_resp = await client.post(
            f"/api/documents/{doc['id']}/pistes",
            json={"texte": "Ajouter un volet RGPD"},
        )
        assert create_resp.status_code == 200
        piste = create_resp.json()
        assert piste["texte"] == "Ajouter un volet RGPD"
        assert piste["status"] == "nouvelle"
        assert piste["document_id"] == doc["id"]

        list_resp = await client.get(f"/api/documents/{doc['id']}/pistes")
        assert list_resp.status_code == 200
        pistes = list_resp.json()
        assert len(pistes) == 1
        assert pistes[0]["id"] == piste["id"]

    @pytest.mark.asyncio
    async def test_create_piste_document_not_found(self, client: AsyncClient):
        """POST .../pistes renvoie 404 si le document n'existe pas."""
        response = await client.post(
            "/api/documents/id-inexistant/pistes",
            json={"texte": "X"},
        )
        assert response.status_code == 404

    @pytest.mark.asyncio
    async def test_create_piste_bumps_document_updated_at(self, client: AsyncClient):
        """Capturer une piste fait remonter le document dans la liste triée
        par updated_at desc."""
        doc = await _create_document(client)
        before = await _document_updated_at(client, doc["id"])

        await client.post(
            f"/api/documents/{doc['id']}/pistes",
            json={"texte": "Une piste qui fait vivre le document"},
        )

        after = await _document_updated_at(client, doc["id"])
        assert after > before

    @pytest.mark.asyncio
    async def test_update_piste_status(self, client: AsyncClient):
        """PATCH /api/documents/pistes/{id} change le statut d'une piste."""
        doc = await _create_document(client)
        create_resp = await client.post(
            f"/api/documents/{doc['id']}/pistes",
            json={"texte": "Piste à explorer"},
        )
        piste_id = create_resp.json()["id"]

        response = await client.patch(
            f"/api/documents/pistes/{piste_id}",
            json={"status": "exploree"},
        )

        assert response.status_code == 200
        assert response.json()["status"] == "exploree"

    @pytest.mark.asyncio
    async def test_update_piste_status_invalid_value(self, client: AsyncClient):
        """PATCH pistes/{id} avec un statut inconnu renvoie 400."""
        doc = await _create_document(client)
        create_resp = await client.post(
            f"/api/documents/{doc['id']}/pistes",
            json={"texte": "Piste à explorer"},
        )
        piste_id = create_resp.json()["id"]

        response = await client.patch(
            f"/api/documents/pistes/{piste_id}",
            json={"status": "statut-bidon"},
        )

        assert response.status_code == 400

    @pytest.mark.asyncio
    async def test_update_piste_status_not_found(self, client: AsyncClient):
        """PATCH pistes/{id} renvoie 404 si la piste n'existe pas."""
        response = await client.patch(
            "/api/documents/pistes/id-inexistant",
            json={"status": "exploree"},
        )
        assert response.status_code == 404


class TestExport:
    """GET /api/documents/{id}/export - réutilise TEL QUEL le circuit des
    documents générés (chat.py::export_conversation) : md écrit dans
    registry.output_dir, docx via registry.execute("docx-pro", ...),
    téléchargement via /api/skills/download/{file_id}."""

    @pytest.mark.asyncio
    async def test_export_markdown_titres_dans_l_ordre_et_telechargement_reel(
        self, client: AsyncClient
    ):
        """3 sections (profondeurs mêlées, ordre non trivial) -> le fichier
        réellement téléchargé contient les titres dans l'ordre `order`,
        avec ## pour depth 0 et ### pour depth 1."""
        doc = await _create_document(client, title="Guide export", brief="Brief du besoin")
        intro = await _create_section(client, doc["id"], "Introduction", order=10.0, depth=0)
        detail = await _create_section(
            client, doc["id"], "Détail bancaire", order=20.0, depth=1
        )
        conclusion = await _create_section(client, doc["id"], "Conclusion", order=30.0, depth=0)
        await client.patch(
            f"/api/documents/sections/{intro['id']}",
            json={"content": "Contenu de l'introduction."},
        )
        await client.patch(
            f"/api/documents/sections/{detail['id']}",
            json={"content": "Contenu du détail bancaire."},
        )
        await client.patch(
            f"/api/documents/sections/{conclusion['id']}",
            json={"content": "Contenu de la conclusion."},
        )

        response = await client.get(f"/api/documents/{doc['id']}/export?format=md")

        assert response.status_code == 200, response.text
        data = response.json()
        assert data["success"] is True
        assert data["format"] == "md"
        assert data["file_name"].endswith(".md")

        # Le fichier est servi par le circuit des documents générés (round-trip réel).
        download = await client.get(data["download_url"])
        assert download.status_code == 200
        text = download.text

        assert text.startswith("# Guide export")
        assert "## Introduction" in text
        assert "### Détail bancaire" in text
        assert "## Conclusion" in text
        assert "Contenu de l'introduction." in text
        assert "Contenu du détail bancaire." in text
        assert "Contenu de la conclusion." in text
        # Ordre respecté : Introduction < Détail bancaire < Conclusion.
        assert text.index("## Introduction") < text.index("### Détail bancaire")
        assert text.index("### Détail bancaire") < text.index("## Conclusion")

    @pytest.mark.asyncio
    async def test_export_section_orpheline_non_vide_en_annexe_pas_dans_le_corps(
        self, client: AsyncClient, db_session
    ):
        """Une section orpheline non vide n'apparaît jamais dans le corps du
        document exporté, mais reste récupérable en annexe (zéro perte)."""
        from app.models.entities import DocumentSection

        doc = await _create_document(client)
        body_section = await _create_section(client, doc["id"], "Corps", order=10.0)
        await client.patch(
            f"/api/documents/sections/{body_section['id']}",
            json={"content": "Contenu du corps."},
        )

        orphan_section = await _create_section(
            client, doc["id"], "Ancienne section détachée", order=20.0
        )
        await client.patch(
            f"/api/documents/sections/{orphan_section['id']}",
            json={"content": "Contenu détaché mais toujours présent."},
        )

        # Aucune route ne pose encore orphan=True (régénération de trame
        # actuelle = remplacement, jamais détachement) - manipulation directe
        # en base, seul moyen de fabriquer ce cas pour l'instant.
        row = await db_session.get(DocumentSection, orphan_section["id"])
        row.orphan = True
        db_session.add(row)
        await db_session.commit()

        response = await client.get(f"/api/documents/{doc['id']}/export?format=md")
        assert response.status_code == 200, response.text

        download = await client.get(response.json()["download_url"])
        text = download.text

        assert "Annexe - sections détachées" in text
        assert "Ancienne section détachée" in text
        assert "Contenu détaché mais toujours présent." in text
        # Pas dans le corps : le titre orphelin n'apparaît qu'APRÈS l'annexe.
        assert text.index("## Corps") < text.index("Annexe - sections détachées")
        assert text.index("Annexe - sections détachées") < text.index(
            "Ancienne section détachée"
        )

    @pytest.mark.asyncio
    async def test_export_document_vide_400(self, client: AsyncClient):
        """Document sans aucune section non-orpheline non vide -> 400."""
        doc = await _create_document(client)
        # Une section existe mais reste vide (statut 'vide', contenu "").
        await _create_section(client, doc["id"], "Section jamais rédigée", order=10.0)

        response = await client.get(f"/api/documents/{doc['id']}/export?format=md")

        assert response.status_code == 400
        assert "rien à exporter" in response.json()["message"].lower()

    @pytest.mark.asyncio
    async def test_export_document_sans_aucune_section_400(self, client: AsyncClient):
        """Document sans la moindre section -> 400 également."""
        doc = await _create_document(client)

        response = await client.get(f"/api/documents/{doc['id']}/export?format=md")

        assert response.status_code == 400

    @pytest.mark.asyncio
    async def test_export_format_inconnu_400(self, client: AsyncClient):
        """Format inconnu -> 400, avant même la résolution du document."""
        response = await client.get("/api/documents/nimporte/export?format=pdf")
        assert response.status_code == 400

    @pytest.mark.asyncio
    async def test_export_document_inexistant_404(self, client: AsyncClient):
        response = await client.get("/api/documents/id-inexistant/export?format=md")
        assert response.status_code == 404

    @pytest.mark.asyncio
    async def test_export_docx_round_trip_reel(self, client: AsyncClient):
        """Le format docx passe par registry.execute("docx-pro", ...). Sans
        bloc de code Python dans le markdown assemblé, CodeGenSkill.execute()
        ne tente AUCUN appel LLM : extract_python_code() ne trouve rien,
        retombe directement sur _fallback_execute() (parser markdown ->
        python-docx, 100% local) - ce chemin est donc exerçable réellement
        en test, sans mock. Le registre de skills n'étant initialisé nulle
        part ailleurs dans cette suite (lifespan de test minimal, cf
        conftest.py - même constat que test_routers_skills.py), on
        l'initialise ici et on le referme en finally pour ne pas polluer le
        reste de la suite (le registre est un singleton de module)."""
        from app.services.skills import close_skills, init_skills

        await init_skills()
        try:
            doc = await _create_document(client, title="Guide export docx")
            section = await _create_section(client, doc["id"], "Corps", order=10.0)
            await client.patch(
                f"/api/documents/sections/{section['id']}",
                json={"content": "Contenu du corps pour le docx."},
            )

            response = await client.get(f"/api/documents/{doc['id']}/export?format=docx")

            assert response.status_code == 200, response.text
            data = response.json()
            assert data["success"] is True
            assert data["format"] == "docx"
            assert data["file_name"].endswith(".docx")

            download = await client.get(data["download_url"])
            assert download.status_code == 200
            assert (
                download.headers["content-type"]
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            assert len(download.content) > 0
        finally:
            await close_skills()


class TestBUG135ExportDocxComplet:
    """BUG-135 - le DOCX exporté reprend TOUT le contenu (une fence de code
    jamais refermée dans une section tronquait le document à partir de là) et
    porte la langue fr-FR. Round-trip réel via l'endpoint + relecture du .docx."""

    @pytest.mark.asyncio
    async def test_export_docx_fence_orpheline_document_complet_et_fr(
        self, client: AsyncClient
    ):
        import io

        from app.services.skills import close_skills, init_skills
        from docx import Document as DocxDocument
        from docx.oxml.ns import qn

        await init_skills()
        try:
            doc = await _create_document(client, title="Rapport complet")
            s1 = await _create_section(client, doc["id"], "Contexte", order=10.0)
            s2 = await _create_section(client, doc["id"], "Conclusion", order=20.0)
            await client.patch(
                f"/api/documents/sections/{s1['id']}",
                json={
                    "content": "Debut du contexte.\n\n```\ngabarit jamais referme\n\nSuite du contexte apres la fence."
                },
            )
            await client.patch(
                f"/api/documents/sections/{s2['id']}",
                json={"content": "Phrase finale de la conclusion."},
            )

            response = await client.get(f"/api/documents/{doc['id']}/export?format=docx")
            assert response.status_code == 200, response.text

            download = await client.get(response.json()["download_url"])
            assert download.status_code == 200

            word = DocxDocument(io.BytesIO(download.content))
            text = "\n".join(p.text for p in word.paragraphs)
            # Zéro perte : tout ce qui suit la fence orpheline est présent.
            assert "Suite du contexte apres la fence." in text
            assert "Phrase finale de la conclusion." in text
            # Titre unique (pas de doublon add_heading + `# titre` du markdown).
            assert len([p for p in word.paragraphs if p.text == "Rapport complet"]) == 1
            # Langue de correction fr-FR posée sur le style Normal.
            normal_rpr = word.styles["Normal"].element.find(qn("w:rPr"))
            lang = normal_rpr.find(qn("w:lang")) if normal_rpr is not None else None
            assert lang is not None and lang.get(qn("w:val")) == "fr-FR"
        finally:
            await close_skills()
