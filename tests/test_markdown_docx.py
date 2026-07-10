"""
BUG-135 - Convertisseur Markdown -> DOCX déterministe (exports Atelier/conversation).

Le chemin historique (registry.execute("docx-pro") -> CodeGenSkill) perdait du
contenu : le strip des fences de code (regex avec alternative `|$` + toggle
in_code_block) effaçait tout le document après une fence ``` jamais refermée,
et pouvait même exécuter un bloc ```python``` légitime du document. Le DOCX
sortait aussi en langue en-US (aucun w:lang posé).

Ces tests spécifient le nouveau module `markdown_docx.render_markdown_docx` :
conversion 100 % locale, zéro perte de contenu, langue fr-FR.
"""

from pathlib import Path

import pytest
from app.services.skills.markdown_docx import render_markdown_docx
from docx import Document
from docx.oxml.ns import qn


def _render(tmp_path: Path, markdown: str) -> Document:
    """Rend le markdown en DOCX puis le rouvre (round-trip disque réel)."""
    output = tmp_path / "out.docx"
    render_markdown_docx(markdown, output)
    assert output.exists() and output.stat().st_size > 0
    return Document(str(output))


def _all_text(doc: Document) -> str:
    """Texte complet du corps (paragraphes + cellules de tableaux)."""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


class TestBUG135ContenuComplet:
    """Le DOCX reprend l'intégralité du markdown, même avec fences piégeuses."""

    def test_fence_orpheline_ne_tronque_pas_la_suite(self, tmp_path: Path):
        """Une fence ``` jamais refermée ne doit plus avaler la fin du
        document (cause racine du BUG-135 : regex `(?:```|$)` + DOTALL)."""
        markdown = (
            "# Rapport\n\n"
            "## Section 1\n\nPremier paragraphe.\n\n"
            "```\ncontenu de bloc jamais referme\n\n"
            "## Section 2\n\nDeuxieme paragraphe.\n\n"
            "## Conclusion\n\nDernier paragraphe du document."
        )
        doc = _render(tmp_path, markdown)
        text = _all_text(doc)
        assert "Premier paragraphe." in text
        assert "Deuxieme paragraphe." in text
        assert "Dernier paragraphe du document." in text

    def test_fence_cloturee_contenu_preserve_en_preformate(self, tmp_path: Path):
        """Un bloc de code clôturé est PRÉSERVÉ (préformaté), pas supprimé :
        un gabarit d'e-mail ou un exemple mis en ``` fait partie du document."""
        markdown = (
            "# Doc\n\n"
            "Avant le bloc.\n\n"
            "```\nBonjour Madame, ceci est un gabarit.\n```\n\n"
            "Apres le bloc."
        )
        doc = _render(tmp_path, markdown)
        text = _all_text(doc)
        assert "Bonjour Madame, ceci est un gabarit." in text
        assert "Avant le bloc." in text
        assert "Apres le bloc." in text

    def test_bloc_python_rendu_pas_execute(self, tmp_path: Path):
        """Un bloc ```python``` du document est du CONTENU : il doit être rendu
        tel quel, jamais extrait/exécuté comme générateur (risque relevé en
        revue : CodeGenSkill.execute exécutait ce code)."""
        # NB sécurité : le os.system ci-dessous est une CHAÎNE INERTE dans un
        # markdown de test - le test garantit précisément qu'elle n'est jamais
        # exécutée (aucun appel shell dans ce test).
        markdown = (
            "# Doc\n\n"
            "```python\nimport os\nos.system('echo pwned')\n```\n\n"
            "Paragraphe final."
        )
        doc = _render(tmp_path, markdown)
        text = _all_text(doc)
        assert "os.system('echo pwned')" in text
        assert "Paragraphe final." in text

    def test_asterisques_non_apparies_preserves(self, tmp_path: Path):
        """L'ancien parser inline jetait les `*` et backticks non appariés."""
        markdown = "# Doc\n\nNote 5* obtenue et cote a 3*2 unites."
        doc = _render(tmp_path, markdown)
        text = _all_text(doc)
        assert "Note 5" in text and "obtenue" in text
        assert "3" in text and "2 unites" in text


class TestBUG135LangueFrancaise:
    """Le DOCX doit porter la langue fr-FR (correction Word), pas en-US."""

    def test_langue_fr_fr_sur_style_normal(self, tmp_path: Path):
        doc = _render(tmp_path, "# Doc\n\nBonjour.")
        normal_rpr = doc.styles["Normal"].element.find(qn("w:rPr"))
        assert normal_rpr is not None
        lang = normal_rpr.find(qn("w:lang"))
        assert lang is not None
        assert lang.get(qn("w:val")) == "fr-FR"

    def test_langue_fr_fr_sur_doc_defaults(self, tmp_path: Path):
        doc = _render(tmp_path, "# Doc\n\nBonjour.")
        langs = doc.styles.element.findall(
            f"{qn('w:docDefaults')}/{qn('w:rPrDefault')}/{qn('w:rPr')}/{qn('w:lang')}"
        )
        assert langs, "w:lang absent des docDefaults"
        assert all(lang.get(qn("w:val")) == "fr-FR" for lang in langs)


class TestBUG135Structure:
    """Titre unique, tableaux réels, constructions markdown courantes."""

    def test_pas_de_doublon_de_titre(self, tmp_path: Path):
        """Le markdown assemblé contient déjà `# titre` : le convertisseur ne
        doit PAS en rajouter un deuxième (défaut relevé en revue : add_heading
        (params.title) + `# titre` du markdown)."""
        doc = _render(tmp_path, "# Mon rapport unique\n\nCorps du texte.")
        occurrences = [p.text for p in doc.paragraphs if p.text == "Mon rapport unique"]
        assert len(occurrences) == 1

    def test_tableau_markdown_devient_vrai_tableau(self, tmp_path: Path):
        markdown = (
            "# Doc\n\n"
            "| Nom | Prix |\n"
            "|-----|------|\n"
            "| FORGER | 490 |\n"
            "| PROPULSER | 2490 |\n"
        )
        doc = _render(tmp_path, markdown)
        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert table.rows[0].cells[0].text == "Nom"
        assert table.rows[0].cells[1].text == "Prix"
        assert table.rows[1].cells[0].text == "FORGER"
        assert table.rows[2].cells[1].text == "2490"

    def test_titre_niveau_4_et_blockquote_rendus(self, tmp_path: Path):
        markdown = (
            "# Doc\n\n"
            "#### Sous-sous-section\n\n"
            "> Une citation importante.\n\n"
            "Fin."
        )
        doc = _render(tmp_path, markdown)
        text = _all_text(doc)
        assert "Sous-sous-section" in text
        assert "#### " not in text, "le #### doit devenir un titre, pas du texte brut"
        assert "Une citation importante." in text
        assert not any(
            p.text.startswith("> ") for p in doc.paragraphs
        ), "le chevron de blockquote ne doit pas rester littéral"

    def test_separateur_horizontal_pas_de_texte_brut(self, tmp_path: Path):
        doc = _render(tmp_path, "# Doc\n\nAvant.\n\n---\n\nApres.")
        text = _all_text(doc)
        assert "Avant." in text and "Apres." in text
        assert not any(p.text.strip() == "---" for p in doc.paragraphs)


@pytest.mark.asyncio
class TestBUG135FallbackLLMDurci:
    """Le fallback historique (DocxSkill._add_content, chemin LLM) garde son
    nettoyage de code résiduel mais ne doit plus tronquer sur fence orpheline."""

    async def test_fallback_fence_orpheline_ne_tronque_plus(self, tmp_path: Path):
        from app.services.skills.docx_generator import DocxSkill

        skill = DocxSkill(output_dir=tmp_path)
        doc = Document()
        skill._add_content(
            doc,
            "## Section 1\nPara 1.\n\n"
            "```python\nprint('code residuel jamais referme')\n\n"
            "## Section 2\nPara 2.\n\n## Conclusion\nFin du document.",
        )
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Para 1." in text
        assert "Fin du document." in text, (
            "une fence orpheline ne doit plus avaler la fin du contenu"
        )

    async def test_fallback_bloc_cloture_toujours_nettoye(self, tmp_path: Path):
        """Comportement voulu du chemin LLM : un bloc ```python``` CLÔTURÉ est
        du code résiduel de génération -> toujours retiré du document."""
        from app.services.skills.docx_generator import DocxSkill

        skill = DocxSkill(output_dir=tmp_path)
        doc = Document()
        skill._add_content(
            doc,
            "Intro.\n\n```python\nprint('residu')\n```\n\nConclusion.",
        )
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "print('residu')" not in text
        assert "Intro." in text and "Conclusion." in text
