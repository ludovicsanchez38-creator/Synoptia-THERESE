"""
Chantier 5 - ExportProfile DOCX (design V2 du 10/07, revue Codex intégrée).

Un profil d'export global (langue, polices, couleurs, footer, marges) pilote
le rendu DOCX déterministe (markdown_docx : exports Atelier + conversations).
Schéma Pydantic versionné strict, JSON dans le data dir en écriture atomique,
fichier corrompu -> repli défauts + avertissement SURFACÉ (pas juste un log),
défauts = charte actuelle (zéro changement si non configuré).
"""
import json
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn
from httpx import AsyncClient


@pytest.fixture()
def clean_profile():
    """Supprime le profil avant ET après chaque test (data dir partagé)."""
    from app.services.export_profile import export_profile_path

    path = export_profile_path()
    path.unlink(missing_ok=True)
    yield path
    path.unlink(missing_ok=True)


class TestExportProfileModel:
    def test_defauts_conformes_a_la_charte_actuelle(self):
        from app.services.export_profile import ExportProfile

        p = ExportProfile()
        assert p.version == 1
        assert p.language == "fr-FR"
        assert p.body_font == "Calibri"
        assert p.body_size_pt == 11
        assert p.heading_font == "Outfit"
        assert "THÉRÈSE" in p.footer_text
        assert p.margins_cm.top == 2.5

    def test_valeurs_bornees(self):
        from app.services.export_profile import ExportProfile
        from pydantic import ValidationError

        with pytest.raises(ValidationError):
            ExportProfile(body_size_pt=500)
        with pytest.raises(ValidationError):
            ExportProfile(heading_color="rouge")
        with pytest.raises(ValidationError):
            ExportProfile(footer_text="x" * 500)
        with pytest.raises(ValidationError):
            ExportProfile(margins_cm={"top": 20, "bottom": 2.5, "left": 2.5, "right": 2.5})


class TestExportProfileStorage:
    def test_absent_defauts_sans_warning(self, clean_profile):
        from app.services.export_profile import load_export_profile

        profile, warning = load_export_profile()
        assert profile.language == "fr-FR"
        assert warning is None

    def test_round_trip_save_load(self, clean_profile):
        from app.services.export_profile import (
            ExportProfile,
            load_export_profile,
            save_export_profile,
        )

        save_export_profile(ExportProfile(language="de-DE", body_font="Arial"))
        profile, warning = load_export_profile()
        assert profile.language == "de-DE"
        assert profile.body_font == "Arial"
        assert warning is None

    def test_corrompu_repli_defauts_warning_fichier_conserve(self, clean_profile):
        from app.services.export_profile import load_export_profile

        clean_profile.write_text("{ pas du json")
        profile, warning = load_export_profile()
        assert profile.language == "fr-FR"  # repli défauts
        assert warning is not None
        # Le fichier fautif n'est PAS écrasé (diagnostic possible).
        assert clean_profile.read_text() == "{ pas du json"

    def test_invalide_repli_defauts_warning(self, clean_profile):
        from app.services.export_profile import load_export_profile

        clean_profile.write_text(json.dumps({"version": 1, "body_size_pt": 999}))
        profile, warning = load_export_profile()
        assert profile.body_size_pt == 11
        assert warning is not None


class TestRenduAvecProfil:
    def _render(self, tmp_path: Path, markdown: str, profile=None) -> Document:
        from app.services.skills.markdown_docx import render_markdown_docx

        out = tmp_path / "out.docx"
        render_markdown_docx(markdown, out, profile=profile)
        return Document(str(out))

    def test_sans_profil_comportement_actuel(self, tmp_path):
        doc = self._render(tmp_path, "# Titre\n\nCorps.")
        normal = doc.styles["Normal"]
        assert normal.font.name == "Calibri"
        lang = normal.element.find(qn("w:rPr")).find(qn("w:lang"))
        assert lang.get(qn("w:val")) == "fr-FR"

    def test_profil_custom_applique(self, tmp_path):
        from app.services.export_profile import ExportProfile

        profile = ExportProfile(
            language="de-DE",
            body_font="Arial",
            body_size_pt=12,
            footer_text="Confidentiel - Ets Toto",
            margins_cm={"top": 3.0, "bottom": 2.0, "left": 2.5, "right": 2.5},
        )
        doc = self._render(tmp_path, "# Titre\n\nCorps du texte.", profile=profile)

        normal = doc.styles["Normal"]
        assert normal.font.name == "Arial"
        assert normal.font.size.pt == 12
        lang = normal.element.find(qn("w:rPr")).find(qn("w:lang"))
        assert lang.get(qn("w:val")) == "de-DE"
        # Footer custom
        footer_text = "".join(
            p.text for p in doc.sections[-1].footer.paragraphs
        )
        assert "Ets Toto" in footer_text
        # Marges (EMU -> cm : 1 cm = 360000 EMU)
        section = doc.sections[0]
        assert abs(section.top_margin.cm - 3.0) < 0.01
        assert abs(section.bottom_margin.cm - 2.0) < 0.01

    def test_couleur_h2_distincte_pilotee(self, tmp_path):
        from app.services.export_profile import ExportProfile

        profile = ExportProfile(h2_color="#FF0000")
        doc = self._render(tmp_path, "# T\n\n## Section\n\nCorps.", profile=profile)
        h2 = doc.styles["Heading 2"]
        assert str(h2.font.color.rgb) == "FF0000"


class TestExportProfileEndpoints:
    @pytest.mark.asyncio
    async def test_get_defauts(self, client: AsyncClient, clean_profile):
        resp = await client.get("/api/config/export-profile")
        assert resp.status_code == 200
        data = resp.json()
        assert data["profile"]["language"] == "fr-FR"
        assert data["warning"] is None

    @pytest.mark.asyncio
    async def test_put_persiste_et_get_relit(self, client: AsyncClient, clean_profile):
        resp = await client.put(
            "/api/config/export-profile",
            json={"version": 1, "language": "en-GB", "footer_text": "Footer perso"},
        )
        assert resp.status_code == 200
        resp = await client.get("/api/config/export-profile")
        assert resp.json()["profile"]["language"] == "en-GB"
        assert resp.json()["profile"]["footer_text"] == "Footer perso"

    @pytest.mark.asyncio
    async def test_put_invalide_422(self, client: AsyncClient, clean_profile):
        resp = await client.put(
            "/api/config/export-profile",
            json={"version": 1, "heading_color": "pas-une-couleur"},
        )
        assert resp.status_code == 422

    @pytest.mark.asyncio
    async def test_delete_reset_defauts(self, client: AsyncClient, clean_profile):
        await client.put(
            "/api/config/export-profile", json={"version": 1, "language": "en-GB"}
        )
        resp = await client.delete("/api/config/export-profile")
        assert resp.status_code == 200
        resp = await client.get("/api/config/export-profile")
        assert resp.json()["profile"]["language"] == "fr-FR"

    @pytest.mark.asyncio
    async def test_export_docx_utilise_le_profil(self, client: AsyncClient, clean_profile):
        """Intégration bout-en-bout : l'export de conversation applique le
        profil sauvegardé (langue + footer lus dans le DOCX rouvert)."""
        import io

        from app.services.skills import close_skills, init_skills

        await client.put(
            "/api/config/export-profile",
            json={"version": 1, "language": "it-IT", "footer_text": "Pied perso"},
        )

        conv = await client.post("/api/chat/conversations", json={"title": "Exp"})
        conv_id = conv.json()["id"]
        # un message pour que l'export ait du contenu
        from app.models.database import get_session_context
        from app.models.entities import Message

        async with get_session_context() as session:
            session.add(Message(conversation_id=conv_id, role="user", content="Bonjour"))
            await session.commit()

        await init_skills()
        try:
            resp = await client.get(f"/api/chat/conversations/{conv_id}/export?format=docx")
            assert resp.status_code == 200, resp.text
            download = await client.get(resp.json()["download_url"])
            word = Document(io.BytesIO(download.content))
        finally:
            await close_skills()

        lang = word.styles["Normal"].element.find(qn("w:rPr")).find(qn("w:lang"))
        assert lang.get(qn("w:val")) == "it-IT"
        footer_text = "".join(p.text for p in word.sections[-1].footer.paragraphs)
        assert "Pied perso" in footer_text
